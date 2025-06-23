from docx import Document
import random
import os
from faker import Faker
from datetime import datetime, timedelta

fake = Faker()

# Ensure output directory exists
output_dir = "doc"
os.makedirs(output_dir, exist_ok=True)

# Example task titles
task_titles = [
    "Develop an Online Booking System",
    "Implement Microservices for Reservation Handling",
    "Build a Secure Payment Processing Service",
    "Optimize Room Availability Management",
    "Enhance User Authentication and Access Control",
    "Integrate a Real-Time Notification System",
    "Design an Interactive Analytics Dashboard",
    "Develop a Mobile Application for Reservations",
    "Improve Database Performance and Scalability",
    "Conduct Comprehensive User Testing and Evaluation"
]

# Example key focus / dynamic aspects
dynamic_aspects = [
    "Defining key success metrics and performance benchmarks.",
    "Identifying potential system vulnerabilities and risk factors.",
    "Analyzing industry trends to incorporate best practices.",
    "Ensuring regulatory compliance and data privacy adherence.",
    "Planning for future scalability and load handling.",
    "Building a prototype for early-stage validation and feedback.",
    "Leveraging AI-driven techniques to enhance performance.",
    "Applying modern CI/CD pipelines for faster deployment.",
    "Monitoring system health with automated observability tools.",
    "Implementing robust authentication mechanisms."
]

# Example additional notes
additional_notes = [
    "Effective teamwork enhances productivity.",
    "Adopting Agile methodologies ensures iterative improvements.",
    "Comprehensive documentation is key to long-term maintainability.",
    "Embracing change is critical for continuous innovation.",
    "Communication with stakeholders is vital for project success.",
    "Continuous testing drives higher software quality.",
    "Regular code reviews foster knowledge sharing.",
    "Monitoring user feedback helps improve usability."
]

# Generate random course code (format: ABCD123)
def generate_course_code():
    letters = ''.join(random.choices('ABCDEFGHIJKLMNOPQRSTUVWXYZ', k=4))
    digits = ''.join(random.choices('0123456789', k=3))
    return f"{letters}{digits}"

# Generate 10 DOCX documents
for i in range(1, 11):
    doc = Document()

    # Generate random course info
    course_code = generate_course_code()
    course_name = f"Advanced {fake.bs().title()}"
    teacher_name = f"Dr. {fake.first_name()} {fake.last_name()}"
    semester = f"Semester {random.choice([1, 2, 3])} {random.choice([2024, 2025])}"
    due_date = datetime.today() + timedelta(days=random.randint(7, 30))
    due_date_str = due_date.strftime("%B %d, %Y")

    # Assignment title
    title = f"{course_code} Assignment {i}: {task_titles[i-1]}"
    doc.add_heading(title, level=1)

    # Course info
    doc.add_paragraph(f"Course Name: {course_name}")
    doc.add_paragraph(f"Instructor: {teacher_name}")
    doc.add_paragraph(f"{semester}")
    doc.add_paragraph(f"Due Date: {due_date_str}")
    doc.add_paragraph(f"Student: {fake.name()} (Student ID: {random.randint(1000000,9999999)})")

    doc.add_paragraph("\n")

    # Purpose of the assignment
    doc.add_heading("Purpose of the Assignment", level=2)
    doc.add_paragraph(fake.paragraph(nb_sentences=5))

    # Assignment requirements
    doc.add_heading("Assignment Requirements", level=2)
    requirements = [
        f"- Understand the core concepts behind {task_titles[i-1]}.",
        f"- Design and implement a functional prototype.",
        "- Document system architecture and workflows.",
        "- Perform user testing and gather feedback.",
        "- Ensure compliance with security best practices.",
        "- Prepare a comprehensive project report.",
        "- Present findings during class presentation sessions.",
        "- Collaborate effectively within a team environment.",
        "- Apply relevant software development methodologies."
    ]
    # Randomize the number of requirements per doc
    for req in random.sample(requirements, k=random.randint(6, len(requirements))):
        doc.add_paragraph(req)

    # Key focus section
    doc.add_heading("Key Focus", level=2)
    doc.add_paragraph(random.choice(dynamic_aspects))

    # Evaluation criteria
    doc.add_heading("Evaluation Criteria", level=2)
    criteria = [
        "- Code Quality and Architecture",
        "- Functionality and Usability",
        "- Innovation and Creativity",
        "- Adherence to Requirements",
        "- Documentation Quality",
        "- Presentation and Communication"
    ]
    for crit in criteria:
        doc.add_paragraph(f"- {crit}")

    # Extra section with random title and paragraph
    doc.add_heading(fake.catch_phrase(), level=2)
    doc.add_paragraph(fake.paragraph(nb_sentences=random.randint(5, 7)))

    # Additional notes section
    doc.add_heading("Additional Notes", level=2)
    doc.add_paragraph(random.choice(additional_notes))

    # Conclusion
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(fake.paragraph(nb_sentences=5))

    # Signature
    doc.add_paragraph(f"\n\nSigned: {fake.name()}")
    doc.add_paragraph(f"Date: {datetime.today().strftime('%B %d, %Y')}")

    # Save the document
    filename = f"{output_dir}/{course_code}_task{i}.docx"
    doc.save(filename)
    print(f"{filename} generated successfully.")
